#!/usr/bin/env python3
"""Generate docs/odoo19-tax-report-upgrade.docx from project documentation."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

OUT = Path(__file__).parent / "odoo19-tax-report-upgrade.docx"


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_title(doc, text):
    doc.add_heading(text, level=0)


def add_h(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows, bold_last_col=False):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)


def add_checklist(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build():
    doc = Document()
    set_doc_defaults(doc)

    add_title(doc, "Odoo 19 Upgrade — Tax Totals on PDF Reports")
    meta = doc.add_paragraph()
    meta.add_run("Module: ").bold = True
    meta.add_run("account (QWeb report template)\n")
    meta.add_run("Template: ").bold = True
    meta.add_run("account.document_tax_totals_template\n")
    meta.add_run("Affected: ").bold = True
    meta.add_run("Sales Orders, Invoices, Purchase Orders\n")
    meta.add_run("Last updated: ").bold = True
    meta.add_run("June 2026")
    doc.add_paragraph()

    add_h(doc, "Table of contents", 1)
    toc = [
        "Executive summary",
        "Symptoms observed",
        "Root causes",
        "Data structure changes (pre-19 vs Odoo 19)",
        "Legacy customizations in this repo",
        "Withholding tax and negative amounts",
        "Fix options",
        "Recommended approach",
        "Implementation checklist",
        "QWeb reference snippet",
        "Testing checklist",
        "Related files",
        "Appendix — Thai summary",
        "Revision history",
    ]
    for i, t in enumerate(toc, 1):
        doc.add_paragraph(f"{i}. {t}")

    add_h(doc, "Executive summary", 1)
    add_p(doc, "After upgrading to Odoo 19, customized PDF tax footers may show:")
    add_bullets(
        doc,
        [
            "VAT amount missing (label visible, value blank)",
            "Wrong tax labels (e.g. Withholding Tax displayed as VAT)",
            "Negative withholding tax (e.g. -5,100.00) on reports",
        ],
    )
    add_p(doc, "These issues are primarily caused by:")
    add_bullets(
        doc,
        [
            "Core API/template changes in _get_tax_totals and account.document_tax_totals_template",
            "Outdated custom QWeb still using pre-19 field names or logic",
            "Misunderstanding WHT sign — negative WHT on Odoo forms is often correct accounting behavior, not a calculation bug",
        ],
    )

    add_h(doc, "Symptoms observed", 1)
    add_table(
        doc,
        ["Symptom", "Example", "Likely cause"],
        [
            ("VAT label without amount", "VAT 7% row empty on quotation PDF", "Amount cell commented out in same_tax_base branch"),
            ("WHT shown as VAT", "ภาษีมูลค่าเพิ่ม/VAT 3% for WHT line", "Hardcoded VAT prefix on all tax_group rows"),
            ("Negative tax line", "Withholding Tax 3%: -5,100.00", "Tax configured to reduce payable; Odoo 19 reports actual sign"),
            ("Totals still correct", "Total = 176,800 with WHT negative", "Backend math OK; display/template issue only"),
            ("Layout unlike pre-upgrade", "Missing/extra rows", "Old loop over groups_by_subtotal no longer valid"),
        ],
    )
    add_h(doc, "Example (correct math, confusing display)", 2)
    add_table(
        doc,
        ["Line", "Amount (THB)"],
        [
            ("Untaxed Amount", "170,000.00"),
            ("Withholding Tax 3%", "-5,100.00"),
            ("VAT 7%", "11,900.00"),
            ("Total", "176,800.00"),
        ],
    )
    add_p(doc, "170,000 - 5,100 + 11,900 = 176,800 (correct)")

    add_h(doc, "Root causes", 1)
    add_h(doc, "1. Odoo 19 redesigned tax totals for reports", 2)
    add_p(
        doc,
        "Odoo centralizes tax breakdown in _get_tax_totals() and renders it via account.document_tax_totals_template.",
    )
    add_bullets(
        doc,
        [
            "Iterates tax_totals['subtotals'] → subtotal['tax_groups']",
            "Uses monetary fields (*_amount_currency) with currency argument",
            "Branches on same_tax_base and display_base_amount_currency",
        ],
    )
    add_p(doc, "Pre-19 custom templates using groups_by_subtotal, formatted_* fields, or manual WHT calculation will not work without migration.")

    add_h(doc, "2. Custom template edits", 2)
    add_bullets(
        doc,
        [
            "Commenting out tax_amount_currency in the same_tax_base branch",
            "Applying one label prefix to every tax group (VAT + WHT merged visually)",
            "Copying old bk_*.xml logic that hardcodes WHT % (e.g. 1%)",
        ],
    )

    add_h(doc, "3. Withholding tax sign (not always a bug)", 2)
    add_bullets(
        doc,
        [
            "VAT 7% — increases total",
            "Withholding Tax 3% — reduces amount due (shown as negative)",
            "Odoo 19 includes WHT in tax_groups with the sign from the tax engine",
        ],
    )

    add_h(doc, "Data structure changes (pre-19 vs Odoo 19)", 1)
    add_table(
        doc,
        ["Pre-19 (typical)", "Odoo 19"],
        [
            ("tax_totals['groups_by_subtotal'][subtotal_name]", "subtotal['tax_groups']"),
            ("group['tax_group_name']", "tax_group['group_name']"),
            ("group['formatted_tax_group_amount']", "tax_group['tax_amount_currency'] + monetary widget"),
            ("subtotal['formatted_amount']", "subtotal['base_amount_currency'] + monetary widget"),
            ("tax_totals['formatted_amount_total']", "tax_totals['total_amount_currency'] + monetary widget"),
            ("Sub-template account.tax_groups_totals", "Inline loop in main template"),
        ],
    )

    add_h(doc, "Important keys in tax_group", 2)
    add_table(
        doc,
        ["Key", "Purpose"],
        [
            ("group_name", "Display name (e.g. VAT 7%, Withholding Tax 3%)"),
            ("tax_amount_currency", "Tax amount in document currency"),
            ("display_base_amount_currency", "Base amount when shown per group"),
            ("same_tax_base (on tax_totals)", "Controls which template branch renders"),
        ],
    )

    add_h(doc, "Legacy customizations in this repo", 1)
    add_table(
        doc,
        ["File", "Behavior"],
        [
            ("bk_amount_tax.xml", "Uses formatted_amount, calls account.tax_groups_totals"),
            ("bk_price_tax.html", "Hides WHT from loop; hardcodes 1% WHT; manual minus sign"),
        ],
    )
    add_p(doc, "Do not deploy these backups on Odoo 19 without migration.", bold=True)

    add_h(doc, "Withholding tax and negative amounts", 1)
    add_h(doc, "Expected behavior (Odoo 19)", 2)
    add_bullets(
        doc,
        [
            "WHT reduces amount due → negative tax_amount_currency on reports is normal",
            "Total still includes WHT algebraically (subtract)",
        ],
    )
    add_h(doc, "Migration implication", 2)
    add_table(
        doc,
        ["Approach", "WHT % source", "Sign on report"],
        [
            ("Odoo 19 standard", "Tax master data", "From engine (often negative)"),
            ("Old custom", "Hardcoded 1%", "Always manual minus"),
        ],
    )

    add_h(doc, "Fix options", 1)
    options = [
        (
            "Option 1 — Inherit Odoo 19 template (recommended)",
            "Restore tax_amount_currency in both branches; conditional VAT/WHT labels; keep monetary widget.",
            "Pros: Correct amounts, maintainable. Cons: Re-test on each upgrade.",
        ),
        (
            "Option 2 — Rename taxes in master data",
            "Accounting → Configuration → Taxes / Tax Groups. Set bilingual names.",
            "Pros: Minimal QWeb change. Cons: WHT may still show negative.",
        ),
        (
            "Option 3 — Accept negative WHT (accounting-aligned)",
            "No abs() on report — match Odoo UI.",
            "Pros: Single source of truth. Cons: Stakeholders may prefer positive deduction display.",
        ),
        (
            "Option 4 — Cosmetic positive WHT on PDF only",
            "Use abs() or separate Deduction label in QWeb only.",
            "Pros: Readable PDF. Cons: Must not confuse with GL/payment amounts.",
        ),
        (
            "Option 5 — Legacy split layout (bk_price_tax.html)",
            "Filter WHT out of loop; dedicated WHT row.",
            "Pros: Full layout control. Cons: High maintenance; wrong % risk.",
        ),
        (
            "Option 6 — Inherit account.tax_groups_totals",
            "Smaller override surface if sub-template still used.",
            "",
        ),
        (
            "Option 7 — Review tax configuration",
            "If wrong on Odoo form: repartition lines, multiple taxes per line, price included.",
            "",
        ),
    ]
    for title, body, note in options:
        add_h(doc, title, 2)
        add_p(doc, body)
        if note:
            add_p(doc, note)

    add_h(doc, "Recommended approach", 1)
    add_bullets(
        doc,
        [
            "Option 1 — Odoo 19 template inherit with full tax_amount_currency in all branches",
            "Option 2 — Correct tax group names in master data",
            "Option 3 — Keep negative WHT unless cosmetic PDF change needed (Option 4)",
            "Avoid Option 5 unless legacy PDF layout is mandatory",
        ],
    )

    add_h(doc, "Implementation checklist", 1)
    add_checklist(
        doc,
        [
            "Identify where custom template lives (Studio, ir.ui.view, or custom module)",
            "Compare against standard Odoo 19 account.document_tax_totals_template",
            "Remove references to groups_by_subtotal, formatted_tax_group_amount, etc.",
            "Add tax_amount_currency monetary span in same_tax_base branch",
            "Replace global VAT prefix with conditional VAT / WHT / default labels",
            "Set tax group names in Accounting configuration",
            "Test SO, Invoice, PO PDFs with VAT-only and VAT+WHT lines",
            "Verify totals match form view (Untaxed, taxes, Total, Amount Due)",
            "Document module version dependency: account (Odoo 19)",
        ],
    )

    add_h(doc, "QWeb reference snippet", 1)
    add_p(doc, "Conditional label (VAT vs WHT vs other):")
    add_code(
        doc,
        """<t t-set="group_name" t-value="tax_group.get('group_name') or ''"/>
<span t-esc="
    'ภาษีมูลค่าเพิ่ม/' + group_name
        if ('vat' in group_name.lower() or 'ภาษีมูลค่าเพิ่ม' in group_name)
    else ('ภาษีหัก ณ ที่จ่าย/' + group_name
        if ('withholding' in group_name.lower() or 'หัก ณ ที่จ่าย' in group_name)
    else group_name)
"/>""",
    )
    add_p(doc, "Amount cell (required in same_tax_base branch):")
    add_code(
        doc,
        """<td class="text-end o_price_total">
    <span class="text-nowrap"
          t-out="tax_group['tax_amount_currency']"
          t-options='{"widget": "monetary", "display_currency": currency}'/>
</td>""",
    )

    add_h(doc, "Testing checklist", 1)
    add_table(
        doc,
        ["Scenario", "Taxes on lines", "Verify"],
        [
            ("VAT only", "7% VAT", "Untaxed, VAT amount, Total"),
            ("VAT + WHT", "7% + 3% WHT", "WHT negative, labels correct, Total"),
            ("Single line high amount", "7% VAT", "Quotation PDF VAT amount visible"),
            ("Paid invoice", "VAT + WHT", "Paid / Amount Due if template extended"),
        ],
    )
    add_p(doc, "Pass criteria: PDF amounts match form; no WHT mislabeled as VAT; Total = Untaxed + taxes + rounding.")

    add_h(doc, "Related files", 1)
    add_table(
        doc,
        ["Path", "Description"],
        [
            ("docs/odoo19-tax-report-upgrade.md", "Markdown source"),
            ("docs/odoo19-tax-report-upgrade.docx", "This document"),
            ("bk_amount_tax.xml", "Pre-19 backup (formatted fields)"),
            ("bk_price_tax.html", "Pre-19 backup (manual WHT 1%)"),
        ],
    )

    add_h(doc, "Appendix — Thai summary", 1)
    add_p(doc, "อาการหลัก: VAT ไม่ขึ้นตัวเลข, WHT โผล่เป็น VAT, WHT ติดลบ")
    add_p(doc, "สาเหตุ: โครง tax_totals เปลี่ยนใน Odoo 19 + template custom ผิด branch / hardcode label")
    add_p(doc, "WHT ติดลบ: มักถูกต้องตามระบบ (ลดยอดจ่าย) ไม่ใช่บั๊กคำนวณ")
    add_p(doc, "แนวทางแก้ที่แนะนำ: inherit template 19 + ตั้งชื่อ tax ใน master + ยอมรับเครื่องหมายลบ WHT (หรือแก้ cosmetic ใน PDF เท่านั้น)")

    add_h(doc, "Revision history", 1)
    add_table(doc, ["Date", "Change"], [("2026-06", "Initial document from upgrade troubleshooting")])

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    build()
